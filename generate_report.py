# -*- coding: utf-8 -*-
"""
Generates VetAI_Diagnostics_FYP_Report.docx (50+ pages) in the SAME layout as the
supplied sample (FYP_REPORT.docx): US Letter, 1" margins, Times New Roman 12,
Heading 1/2/3 = 18/16/14 bold, two title pages, declaration/preface/etc., TOC,
lists, nomenclature, Chapters 1-6, References, Appendices.

Content is grounded in the real FinalFYP (VetAI Diagnostics) codebase and written
in a natural, project-specific academic voice.
"""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# ----------------------- global formatting -----------------------
sec = doc.sections[0]
sec.page_width = Inches(8.5)
sec.page_height = Inches(11)
for m in ("top_margin", "bottom_margin", "left_margin", "right_margin"):
    setattr(sec, m, Inches(1))

normal = doc.styles["Normal"]
normal.font.name = "Times New Roman"
normal.font.size = Pt(12)
normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
pf = normal.paragraph_format
pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
pf.space_after = Pt(6)

for name, size in (("Heading 1", 18), ("Heading 2", 16), ("Heading 3", 14)):
    st = doc.styles[name]
    st.font.name = "Times New Roman"
    st.font.size = Pt(size)
    st.font.bold = True
    st.font.color.rgb = RGBColor(0x1F, 0x1F, 0x1F)
    st.paragraph_format.space_before = Pt(12)
    st.paragraph_format.space_after = Pt(6)


# ----------------------- helpers -----------------------
def para(text="", align=None, bold=False, italic=False, size=None, space_after=None):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    if text:
        r = p.add_run(text)
        r.bold = bold
        r.italic = italic
        if size:
            r.font.size = Pt(size)
    if space_after is not None:
        p.paragraph_format.space_after = Pt(space_after)
    return p


def body(text):
    p = doc.add_paragraph(text)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    return p


def h1(text):
    return doc.add_heading(text, level=1)


def h2(text):
    return doc.add_heading(text, level=2)


def h3(text):
    return doc.add_heading(text, level=3)


def bullet(text):
    return doc.add_paragraph(text, style="List Bullet")


def numbered(text):
    return doc.add_paragraph(text, style="List Number")


def page_break():
    doc.add_page_break()


def figure_placeholder(caption):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("[ Insert figure / diagram / screenshot here ]")
    r.italic = True
    r.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    c = doc.add_paragraph()
    c.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rc = c.add_run(caption)
    rc.bold = True
    rc.font.size = Pt(11)


def caption(text):
    c = doc.add_paragraph()
    c.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rc = c.add_run(text)
    rc.bold = True
    rc.font.size = Pt(11)


def make_table(headers, rows, widths=None, title=None):
    if title:
        caption(title)
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = t.rows[0].cells
    for i, htxt in enumerate(headers):
        hdr[i].text = ""
        rp = hdr[i].paragraphs[0].add_run(htxt)
        rp.bold = True
        rp.font.size = Pt(11)
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), "1F3864")
        hdr[i]._tc.get_or_add_tcPr().append(shd)
        rp.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            rr = cells[i].paragraphs[0].add_run(str(val))
            rr.font.size = Pt(10.5)
    if widths:
        for i, w in enumerate(widths):
            for row in t.rows:
                row.cells[i].width = Inches(w)
    doc.add_paragraph()
    return t


def add_toc():
    p = doc.add_paragraph()
    run = p.add_run()
    fldBegin = OxmlElement("w:fldChar"); fldBegin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText"); instr.set(qn("xml:space"), "preserve")
    instr.text = 'TOC \\o "1-3" \\h \\z \\u'
    fldSep = OxmlElement("w:fldChar"); fldSep.set(qn("w:fldCharType"), "separate")
    t = OxmlElement("w:t"); t.text = "Right-click here and choose 'Update Field' to build the Table of Contents."
    fldEnd = OxmlElement("w:fldChar"); fldEnd.set(qn("w:fldCharType"), "end")
    for el in (fldBegin, instr, fldSep, t, fldEnd):
        run._r.append(el)


def code_block(lines):
    for ln in lines:
        p = doc.add_paragraph()
        r = p.add_run(ln)
        r.font.name = "Consolas"
        r.font.size = Pt(9.5)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    doc.add_paragraph()


CENTER = WD_ALIGN_PARAGRAPH.CENTER
TITLE = "VetAI Diagnostics: A Deep Learning and Generative AI System for Cattle Disease Detection and Diagnosis"

# =====================================================================
# TITLE PAGE 1
# =====================================================================
para(); para()
para(TITLE, align=CENTER, bold=True, size=18)
for _ in range(6):
    para()
figure_placeholder("University / Department Logo")
for _ in range(5):
    para()
para("DEPARTMENT OF TECHNOLOGY", align=CENTER, bold=True, size=14)
para("THE UNIVERSITY OF LAHORE", align=CENTER, bold=True, size=14)
para("Spring, 2026", align=CENTER, bold=True, size=14)
page_break()

# =====================================================================
# TITLE PAGE 2
# =====================================================================
para()
para(TITLE, align=CENTER, bold=True, size=15)
para(); para()
para("A thesis submitted in partial fulfilment of the requirements for the degree of "
     "Bachelor of Science in Information Engineering Technology at The University of Lahore, Pakistan.",
     align=CENTER)
para(); para()
tb = doc.add_table(rows=2, cols=2); tb.style = "Table Grid"; tb.alignment = WD_TABLE_ALIGNMENT.CENTER
tb.rows[0].cells[0].paragraphs[0].add_run("Submitted By:").bold = True
tb.rows[0].cells[1].paragraphs[0].add_run("Submitted By:").bold = True
tb.rows[1].cells[0].text = "[Student Name]"
tb.rows[1].cells[1].text = "[Roll No. / Registration No.]"
para()
ts = doc.add_table(rows=2, cols=2); ts.style = "Table Grid"; ts.alignment = WD_TABLE_ALIGNMENT.CENTER
ts.rows[0].cells[0].paragraphs[0].add_run("Supervised By:").bold = True
ts.rows[0].cells[1].paragraphs[0].add_run("Supervised By:").bold = True
ts.rows[1].cells[0].text = "Supervisor"
ts.rows[1].cells[1].text = "[Supervisor Name]"
para(); para(); para(); para()
para("DEPARTMENT OF TECHNOLOGY", align=CENTER, bold=True, size=14)
para("THE UNIVERSITY OF LAHORE", align=CENTER, bold=True, size=14)
para("Spring, 2026", align=CENTER, bold=True, size=12)
page_break()

# =====================================================================
# CERTIFICATE OF APPROVAL
# =====================================================================
h1("Certificate of Approval")
body("This is to certify that the work presented in this thesis, titled “VetAI Diagnostics: A Deep "
     "Learning and Generative AI System for Cattle Disease Detection and Diagnosis,” was carried out by "
     "the undersigned student under my supervision and is approved for submission to the Department of "
     "Technology, The University of Lahore, in partial fulfilment of the requirements for the degree of "
     "Bachelor of Science in Information Engineering Technology.")
para()
sig = doc.add_table(rows=3, cols=2); sig.style = "Table Grid"; sig.alignment = WD_TABLE_ALIGNMENT.CENTER
cells = [
    ("Supervisor\nName: [Supervisor Name]\nSignature: ___________________",
     "Evaluator 1\nName: ___________________\nSignature: ___________________"),
    ("Evaluator 2\nName: ___________________\nSignature: ___________________",
     "Convener, Project Cell\nName: ___________________\nSignature: ___________________"),
    ("Head, Department of Technology\nSignature: ___________________________",
     "Date: ___________________________"),
]
for r, (a, b) in enumerate(cells):
    sig.rows[r].cells[0].text = a
    sig.rows[r].cells[1].text = b
page_break()

# =====================================================================
# DECLARATION
# =====================================================================
h1("Declaration")
body("I declare that the contents of this thesis are entirely my own work, except where the contribution "
     "of others has been explicitly acknowledged and cited. The software described here—the VetAI "
     "Diagnostics platform, including its Django backend, React frontend, the convolutional neural network "
     "used for image classification, the symptom-matching engine, and the generative-AI integration layer"
     "—was designed and implemented by me as part of my final year project.")
body("No part of this thesis has been or will be submitted for any other degree or examination at this "
     "university or elsewhere. Where I have used external libraries, pre-trained models, public datasets, "
     "or third-party services such as the Google Gemini API, these have been clearly identified in the "
     "relevant sections of the document.")
body("This thesis is protected by copyright, and no part of it may be reproduced, quoted, or published "
     "without the express written consent of the author and the supervisor.")
para()
para("Student Name: ______________________________________________")
para()
para("Signature: _________________________________________________")
para()
para("Date: _____________________________________________________")
page_break()

# =====================================================================
# PREFACE
# =====================================================================
h1("Preface")
body("This work was carried out in the Department of Technology, The University of Lahore, between September "
     "2025 and April 2026. It brings together two areas that interested me throughout my degree: computer "
     "vision and the practical application of large language models.")
body("The idea grew out of a simple observation. In many rural areas, livestock owners do not have quick "
     "access to a veterinarian, and a delay of even a day or two in identifying a disease such as Foot and "
     "Mouth or Lumpy Skin Disease can be costly. I wanted to build something that could give a farmer a fast, "
     "first-level indication of what might be wrong—from a photograph or from a short list of observed "
     "symptoms—while making it very clear that the tool supports, and does not replace, professional "
     "veterinary advice.")
body("The result is VetAI Diagnostics, a web application that combines a trained image classifier with a "
     "symptom-based reasoning engine and a generative AI layer that produces up-to-date guidance on "
     "treatment and prevention. The chapters that follow describe how the system was designed, the choices "
     "made during implementation, the problems encountered along the way, and the results obtained during "
     "testing. I have tried to be honest about what works well and what does not, because I believe an "
     "accurate account of the limitations is more useful to a future reader than an overstated claim of "
     "success.")
page_break()

# =====================================================================
# ACKNOWLEDGEMENT
# =====================================================================
h1("Acknowledgement")
body("I am grateful to my supervisor, [Supervisor Name], whose guidance kept the project on track and whose "
     "questions repeatedly pushed me to justify my design decisions rather than accept the first thing that "
     "happened to work.")
body("I would like to thank the Department of Technology at The University of Lahore for providing the "
     "environment, the laboratory access, and the resources needed to complete this work.")
body("A special mention goes to Google for access to the Gemini API, which made the real-time medical "
     "information layer possible, and to the open-source community behind TensorFlow, Django, and React, "
     "whose documentation and discussion forums I relied on more times than I can count.")
body("Finally, I thank my family for their patience and support during the long stretches of debugging, "
     "particularly the days spent getting the trained model to load reliably across different TensorFlow "
     "versions, and my classmates for testing early versions of the application and pointing out the rough "
     "edges I had stopped noticing.")
page_break()

# =====================================================================
# DETAILS OF COLLABORATION
# =====================================================================
h1("Details of Collaboration and External Resources")
h2("Third-Party APIs and Services")
bullet("Google Gemini API (gemini-2.5-flash): used to generate structured, real-time disease information "
       "— symptoms, severity, treatment, prevention, recommended antibiotics, and contagion status "
       "— and to power the conversational Vet AI Chat assistant.")
bullet("Railway: cloud platform used to host the Django backend in production.")
bullet("Vercel: platform used to host the React (Vite) frontend.")
h2("Open-Source Libraries and Frameworks")
bullet("Django and Django REST Framework (backend web framework and REST API layer).")
bullet("djangorestframework-simplejwt (JWT-based authentication).")
bullet("TensorFlow and Keras with the MobileNetV2 backbone (deep learning model and transfer learning).")
bullet("scikit-learn, NumPy, and pandas (metrics, numerical processing, and dataset handling).")
bullet("Pillow and OpenCV (image loading and processing).")
bullet("React 19 with Vite, Tailwind CSS, TanStack React Query, Axios, Recharts, and lucide-react (frontend).")
bullet("MySQL with the mysqlclient driver (relational database for users and diagnosis records).")
bullet("Gunicorn and WhiteNoise (production server and static-file serving).")
h2("Datasets")
bullet("A labelled image dataset organised by class (Foot and Mouth Disease, Lumpy Skin Disease, healthy), "
       "used to train and validate the convolutional neural network.")
bullet("A tabular symptom dataset (Training.csv and Testing.csv) of roughly two thousand records and about "
       "ninety-three symptom indicators, used by the symptom-matching engine.")
page_break()

# =====================================================================
# ABSTRACT
# =====================================================================
h1("Abstract")
body("Timely identification of disease in cattle is a recurring problem for small and medium livestock "
     "owners, who often cannot reach a veterinarian quickly. This thesis presents VetAI Diagnostics, a "
     "full-stack web application that provides a first-level diagnostic aid through two complementary paths. "
     "The first path accepts a photograph of an animal and classifies it using a convolutional neural "
     "network built on MobileNetV2 with transfer learning; the trained model distinguishes between Foot and "
     "Mouth Disease, Lumpy Skin Disease, and a healthy animal, and reports a confidence score for each "
     "prediction. The second path takes a set of observed symptoms selected by the user and ranks the most "
     "likely diseases using a transparent scoring engine built from a tabular dataset of roughly two thousand "
     "records and ninety-three symptom indicators.")
body("For both paths, the system queries the Google Gemini API in real time to obtain current medical "
     "guidance, including treatment steps, prevention measures, recommended antibiotics, and whether the "
     "condition is contagious. A local knowledge base acts as a fallback when the API is unavailable. The "
     "application also includes a veterinary chat assistant, a per-user analytics dashboard, and a complete "
     "history of past diagnoses. The backend is implemented in Django and Django REST Framework with JSON "
     "Web Token authentication, while the frontend is a React single-page application built with Vite and "
     "styled with Tailwind CSS.")
body("During development, particular attention was paid to the reliability of the generative layer. The "
     "Gemini integration was hardened so that empty or safety-filtered responses no longer cause a silent "
     "failure: the service now validates that a response actually contains usable text, retries across "
     "multiple model versions, and relaxes overly strict content filters that were blocking legitimate "
     "veterinary text. This thesis describes the architecture, the implementation, the testing carried out, "
     "and the limitations that remain, and it outlines a path for extending the system to more diseases and "
     "additional animal species. The work demonstrates that an inexpensive, accessible diagnostic aid can be "
     "assembled from widely available components, provided the generative layer is engineered defensively "
     "and the limits of the tool are communicated honestly to the user.")
page_break()

# =====================================================================
# TABLE OF CONTENTS
# =====================================================================
h1("Table of Contents")
add_toc()
page_break()

# =====================================================================
# LIST OF FIGURES
# =====================================================================
h1("List of Figures")
make_table(
    ["Figure No.", "Description", "Page"],
    [
        ["Figure 1.1", "Manual veterinary diagnosis versus the proposed AI-assisted workflow", ""],
        ["Figure 1.2", "Iterative development methodology followed in the project", ""],
        ["Figure 3.1", "Overall three-tier system architecture", ""],
        ["Figure 3.2", "Use-case diagram for the registered user", ""],
        ["Figure 3.3", "Image diagnosis pipeline (upload to stored report)", ""],
        ["Figure 3.4", "Symptom diagnosis scoring and enrichment flow", ""],
        ["Figure 3.5", "Gemini request, retry, and fallback logic", ""],
        ["Figure 3.6", "Entity-relationship diagram of the database", ""],
        ["Figure 4.1", "Backend application structure (accounts, animal, ml_model)", ""],
        ["Figure 4.2", "Convolutional neural network architecture", ""],
        ["Figure 4.3", "Frontend component hierarchy", ""],
        ["Figure 4.4", "JWT authentication token flow", ""],
        ["Figure 5.1", "Confidence distribution across sample predictions", ""],
        ["Figure 5.2", "Application screenshots (dashboard, diagnosis, report)", ""],
    ],
    widths=[1.2, 4.6, 0.8],
)
page_break()

# =====================================================================
# LIST OF TABLES
# =====================================================================
h1("List of Tables")
make_table(
    ["Table No.", "Description", "Page"],
    [
        ["Table 1.1", "Functional and non-functional requirements", ""],
        ["Table 3.1", "Primary use cases", ""],
        ["Table 4.1", "Technology stack used in the project", ""],
        ["Table 4.2", "CNN layer configuration", ""],
        ["Table 4.3", "REST API endpoints", ""],
        ["Table 4.4", "Frontend components and responsibilities", ""],
        ["Table 5.1", "Trained model classes and supported diseases", ""],
        ["Table 5.2", "Functional test cases and outcomes", ""],
        ["Table 5.3", "Gemini reliability improvements", ""],
        ["Table 5.4", "Comparison with existing approaches", ""],
        ["Table B.1", "Environment variables reference", ""],
        ["Table C.1", "AnimalDetection database fields", ""],
        ["Table C.2", "SymptomDiagnosis database fields", ""],
    ],
    widths=[1.2, 4.6, 0.8],
)
page_break()

# =====================================================================
# NOMENCLATURE
# =====================================================================
h1("Important Nomenclature")
h2("Symbols and Notation")
make_table(
    ["Symbol", "Description"],
    [
        ["n", "Number of disease classes the model can predict"],
        ["c", "Confidence score (softmax probability) of a prediction, in the range 0 to 1"],
        ["m", "Number of symptoms matched against a candidate disease"],
        ["match_ratio", "Matched symptoms divided by the number of symptoms the user selected"],
        ["coverage", "Matched symptoms divided by the total symptoms recorded for a disease"],
        ["w_i", "Class weight applied to class i to compensate for class imbalance during training"],
    ],
    widths=[1.6, 5.0],
)
h2("Abbreviations")
make_table(
    ["Abbreviation", "Full Form"],
    [
        ["AI", "Artificial Intelligence"], ["API", "Application Programming Interface"],
        ["ASGI", "Asynchronous Server Gateway Interface"], ["CNN", "Convolutional Neural Network"],
        ["CORS", "Cross-Origin Resource Sharing"], ["CSV", "Comma-Separated Values"],
        ["DRF", "Django REST Framework"], ["FMD", "Foot and Mouth Disease"],
        ["GPU", "Graphics Processing Unit"], ["HTTP", "Hypertext Transfer Protocol"],
        ["JSON", "JavaScript Object Notation"], ["JWT", "JSON Web Token"],
        ["LLM", "Large Language Model"], ["LSD", "Lumpy Skin Disease"],
        ["ML", "Machine Learning"], ["ORM", "Object-Relational Mapping"],
        ["REST", "Representational State Transfer"], ["SPA", "Single-Page Application"],
        ["SQL", "Structured Query Language"], ["UI", "User Interface"],
        ["WSGI", "Web Server Gateway Interface"],
    ],
    widths=[2.0, 4.6],
)
page_break()

# =====================================================================
# CHAPTER 1: INTRODUCTION
# =====================================================================
h1("Chapter 1: Introduction")
h2("1.1 Motivation")
body("Livestock is a major part of the rural economy in Pakistan, and the health of cattle directly affects "
     "the income of millions of households. When an animal falls ill, the speed of the response matters. "
     "Diseases such as Foot and Mouth Disease and Lumpy Skin Disease spread quickly within a herd, and an "
     "owner who can recognise the early signs and isolate the affected animal stands a far better chance of "
     "limiting the damage. The difficulty is that qualified veterinarians are not always nearby, and a farm "
     "visit can take time to arrange, especially in remote districts where a single veterinary officer may be "
     "responsible for a very large area.")
body("Smartphones, on the other hand, are widely available even in rural areas, and mobile data coverage has "
     "improved considerably in recent years. This creates an opportunity. If a phone can be used to take a "
     "photograph or to answer a few simple questions about what an animal is showing, then a software system "
     "could provide an immediate, preliminary indication of the problem. Such a tool would not replace a "
     "veterinarian, but it could help an owner decide how urgently to seek professional help and what to do "
     "in the meantime, such as separating a possibly contagious animal from the rest of the herd.")
body("This project was motivated by the wish to build exactly that kind of tool, and to do so using "
     "techniques that have matured in recent years—image classification with deep neural networks, and "
     "text generation with large language models—combined in a way that is genuinely useful rather than "
     "a technology demonstration. The intention throughout was to keep the system practical: it should run on "
     "modest hardware, work over an ordinary internet connection, and present its results in plain language.")

h2("1.2 Background and Context")
body("Foot and Mouth Disease is a highly contagious viral infection of cloven-hoofed animals. It rarely "
     "kills adult cattle outright, but it causes fever, blisters, lameness, and a sharp drop in milk "
     "production, and an outbreak can spread through a herd and to neighbouring farms with alarming speed. "
     "Lumpy Skin Disease, also viral and spread mainly by biting insects, produces the characteristic skin "
     "nodules that give the disease its name, along with fever and weakness. Both diseases have a visible "
     "presentation, which is precisely what makes them suitable candidates for image-based detection.")
body("Beyond these two conditions, cattle suffer from a wide range of ailments—mastitis, bloat, "
     "blackleg, pneumonia, and many others—whose signs are not always obvious from a single "
     "photograph but which can often be narrowed down from a description of the symptoms. This dual nature of "
     "the problem, where some conditions are best identified visually and others from reported signs, "
     "directly shaped the decision to offer two diagnostic paths rather than one.")
body("The recent availability of capable, affordable generative models adds a third dimension. Maintaining a "
     "hand-written medical database that stays current is laborious. A generative model, properly "
     "constrained, can supply up-to-date guidance on demand, which keeps the application's medical content "
     "fresh without a large editorial effort. The challenge, addressed at length in this thesis, is making "
     "that generative layer reliable enough to depend on.")

h2("1.3 Problem Statement")
body("There is no shortage of general information about cattle diseases online, but it is scattered, often "
     "written for specialists, and not tied to what a particular owner is actually observing. A farmer with "
     "a sick animal needs three things quickly: a likely name for the condition, an indication of how serious "
     "it is, and a short list of sensible next steps. Existing resources rarely deliver all three in one "
     "place, and almost none of them accept a photograph as input or adapt their answer to the symptoms a "
     "specific animal is showing.")
body("The problem this project addresses is therefore the absence of an accessible, low-cost system that can "
     "take everyday inputs—an image of an animal, or a set of observed symptoms—and return a "
     "structured, readable assessment that includes a probable diagnosis, a severity level, treatment and "
     "prevention guidance, and a clear warning when a condition is contagious, all while reminding the user "
     "that the result must be confirmed by a professional.")

h2("1.4 Aims and Objectives")
h3("1.4.1 Primary Aim")
body("To design and build a web-based diagnostic aid that identifies common cattle diseases from images and "
     "from reported symptoms, and that presents trustworthy, real-time medical guidance for each result.")
h3("1.4.2 Specific Objectives")
numbered("Train a convolutional neural network, using transfer learning on MobileNetV2, to classify cattle "
         "images into Foot and Mouth Disease, Lumpy Skin Disease, and healthy.")
numbered("Build a preprocessing and inference pipeline that loads the trained model reliably and returns a "
         "prediction with a confidence score for each class.")
numbered("Implement a symptom-based diagnosis engine that scores candidate diseases from a tabular dataset "
         "of symptoms and returns the most likely matches with the symptoms that were matched.")
numbered("Integrate the Google Gemini API to enrich every diagnosis with current treatment, prevention, "
         "and medication guidance, with a local fallback when the API cannot be reached.")
numbered("Provide a conversational veterinary assistant through a secure server-side proxy so that the API "
         "key is never exposed to the browser.")
numbered("Implement user registration, JWT-based authentication, and per-user storage of diagnosis history.")
numbered("Build an analytics dashboard that summarises a user's detections with charts and key statistics.")
numbered("Deliver the whole system as a responsive single-page application backed by a documented REST API, "
         "and deploy it to a public hosting platform.")

h2("1.5 Significance of the Study")
body("The value of this work is partly practical and partly methodological. Practically, it produces a "
     "working tool that could help livestock owners act sooner when an animal shows signs of illness, which "
     "in turn can reduce losses and limit the spread of contagious disease. Methodologically, it shows how "
     "three quite different forms of artificial intelligence—a trained image classifier, an "
     "interpretable rule-based engine, and a generative language model—can be combined within a single "
     "conventional web application, and it documents the engineering needed to make the most fragile of those "
     "components, the generative layer, behave dependably in production.")

h2("1.6 Scope of the Project")
h3("1.6.1 In Scope")
bullet("Image-based classification of three cattle conditions (Foot and Mouth, Lumpy Skin, healthy).")
bullet("Symptom-based diagnosis using a curated dataset and a transparent scoring method.")
bullet("Real-time generation of disease guidance through the Gemini API, with a JSON fallback.")
bullet("A veterinary question-and-answer chat assistant.")
bullet("User accounts, authentication, and persistent diagnosis history.")
bullet("A statistics dashboard and a detailed, printable-style diagnosis report.")
h3("1.6.2 Out of Scope")
bullet("Diagnosis of species other than cattle.")
bullet("Real-time video analysis or continuous monitoring.")
bullet("Integration with government livestock databases or veterinary clinic systems.")
bullet("Native mobile applications for Android or iOS.")
bullet("Prescription or dosing decisions, which remain the responsibility of a licensed veterinarian.")

h2("1.7 Development Methodology")
body("The project followed an iterative, incremental approach rather than a single waterfall pass. Work was "
     "organised into short cycles, each of which delivered a working slice of the system: first the "
     "authentication and the basic application shell, then the symptom engine, then the image model and its "
     "inference pipeline, then the Gemini integration, and finally the dashboard and history features. Each "
     "cycle ended with manual testing, after which the next set of features was planned. This style suited a "
     "single-developer project with an evolving understanding of the problem, and it meant that a usable "
     "version existed at every stage rather than only at the end.")
figure_placeholder("Figure 1.2: Iterative development methodology followed in the project")
make_table(
    ["ID", "Requirement", "Type"],
    [
        ["FR-1", "A user can register and log in securely", "Functional"],
        ["FR-2", "A user can upload an image and receive a disease prediction", "Functional"],
        ["FR-3", "A user can select symptoms and receive ranked candidate diseases", "Functional"],
        ["FR-4", "Each diagnosis includes treatment, prevention, and contagion guidance", "Functional"],
        ["FR-5", "A user can view a history of past diagnoses and summary statistics", "Functional"],
        ["FR-6", "A user can ask veterinary questions through a chat assistant", "Functional"],
        ["NFR-1", "Predictions complete within a few seconds under normal load", "Non-functional"],
        ["NFR-2", "The API key and database are never exposed to the client", "Non-functional"],
        ["NFR-3", "The interface is responsive on both desktop and mobile screens", "Non-functional"],
        ["NFR-4", "The system degrades gracefully when the Gemini API is unavailable", "Non-functional"],
    ],
    widths=[0.8, 4.4, 1.4],
    title="Table 1.1: Functional and non-functional requirements",
)

h2("1.8 Organization of the Dissertation")
body("The remainder of this dissertation is organised as follows. Chapter 2 reviews related work in image "
     "classification for animal health, transfer learning, symptom-based expert systems, and the use of "
     "large language models for domain information, and it identifies the gap the project fills. Chapter 3 "
     "describes the overall architecture, the requirements, the use cases, and the design of each major "
     "component. Chapter 4 covers the implementation in detail, including the technology stack, the backend "
     "services, the model training, the frontend, the authentication system, the API, and the deployment. "
     "Chapter 5 presents the results and a discussion of the system's behaviour, including testing, the work "
     "done to make the generative layer reliable, a security assessment, and a comparison with existing "
     "solutions. Chapter 6 draws conclusions, states the limitations honestly, and suggests directions for "
     "future work.")
page_break()

# =====================================================================
# CHAPTER 2: LITERATURE REVIEW
# =====================================================================
h1("Chapter 2: Literature Review")
h2("2.1 Deep Learning for Image Classification")
body("The accuracy of image classification improved dramatically once deep convolutional networks were "
     "trained on large datasets such as ImageNet. Early architectures like AlexNet showed that a deep "
     "network trained on enough labelled images could outperform hand-engineered features by a wide margin. "
     "Later networks such as VGG and GoogLeNet went deeper and more carefully structured, and ResNet "
     "introduced residual connections that allowed very deep networks to be trained without the gradients "
     "vanishing. For this project, the relevant question was not which network achieves the highest absolute "
     "accuracy, but which one is small and fast enough to run on a modest server while still giving usable "
     "results on a limited dataset.")

h2("2.2 Convolutional Neural Networks")
body("A convolutional neural network learns a hierarchy of features directly from pixels. The early layers "
     "respond to simple patterns such as edges and colour gradients; deeper layers combine these into "
     "textures, parts, and eventually whole objects. Three ideas make this efficient: local receptive "
     "fields, so each neuron looks at a small region; weight sharing, so the same filter is applied across "
     "the whole image; and pooling, which reduces spatial resolution while keeping the most salient "
     "information. These properties make convolutional networks particularly well suited to the kind of "
     "visual signs—skin nodules, blisters, swelling—that distinguish the diseases targeted in this "
     "work.")

h2("2.3 Transfer Learning")
body("Training a deep network from scratch needs a large, labelled dataset that is rarely available for a "
     "specialised problem like cattle disease. Transfer learning sidesteps this by reusing a network already "
     "trained on a general dataset and adapting only the final layers to the new task. In practice the "
     "convolutional base is frozen so that its learned features are preserved, and a small classification "
     "head is trained on top. Because the base has already learned generally useful visual features, far "
     "fewer task-specific images are needed to reach good performance, and training is much faster. This is "
     "the approach taken here: the MobileNetV2 base keeps its ImageNet weights, and only the dense layers "
     "added above it are trained on the cattle images.")

h2("2.4 The MobileNet Family")
body("MobileNetV2, introduced by Sandler and colleagues, was designed for environments where compute is "
     "constrained, such as mobile phones and small servers. Its central innovation is the inverted residual "
     "block with a linear bottleneck, built on depthwise separable convolutions. A depthwise separable "
     "convolution factorises a standard convolution into a per-channel spatial filter followed by a "
     "pointwise combination, which drastically reduces the number of multiplications and parameters while "
     "preserving most of the representational power. This efficiency is the main reason MobileNetV2 was "
     "chosen as the backbone for the present work over heavier networks such as ResNet-50 or Inception.")

h2("2.5 Computer Vision for Animal and Plant Disease")
body("Image-based disease detection has been applied widely in agriculture, most successfully in plant "
     "pathology, where convolutional networks trained on leaf images can identify crop diseases with high "
     "accuracy. The same pattern—transfer learning on a pre-trained backbone, fine-tuned on a "
     "domain-specific image set—has more recently been applied to animal health, including the "
     "detection of skin conditions and external signs of disease in livestock. The literature consistently "
     "reports that the quality and balance of the training images matters more than the exact choice of "
     "architecture, which informed the use of class weighting and data augmentation during training in this "
     "project.")

h2("2.6 Symptom-Based and Expert Systems")
body("Before machine learning became widespread, medical decision support relied on rule-based expert "
     "systems that matched reported symptoms against known disease profiles. Classic systems of this kind "
     "encoded expert knowledge as explicit rules and could explain their reasoning, which is valuable in a "
     "setting where a user may want to understand why a particular disease was suggested. The symptom engine "
     "in this project follows that tradition rather than training an opaque classifier: it uses an "
     "interpretable scoring formula over a symptom dataset, which keeps the reasoning visible and the "
     "behaviour predictable, and which does not require a large labelled training set of symptom-to-disease "
     "examples.")

h2("2.7 Large Language Models for Domain Information")
body("Large language models such as Google's Gemini family can produce fluent, structured text on demand, "
     "which makes them attractive for generating up-to-date guidance instead of maintaining a large static "
     "database by hand. The trade-off is that their output must be constrained and validated. A model asked "
     "for free-form text may wander, hallucinate, or return content that is hard to parse. The common "
     "remedy, adopted in this work, is to instruct the model to return strictly formatted JSON, to request a "
     "JSON response type, and to validate and repair the result before it is shown to the user.")

h2("2.8 Prompt Engineering and Structured Output")
body("Getting reliable, machine-readable output from a language model is a discipline in its own right. "
     "Effective prompts state the required format explicitly, give an example of that format, constrain the "
     "length of each field, and forbid extraneous text such as markdown fences or explanations. Lowering the "
     "sampling temperature makes the output more deterministic, which is desirable when the goal is a "
     "structured record rather than a creative passage. A further practical issue is that safety filters can "
     "block legitimate content: veterinary text that mentions antibiotics, severe illness, or the death of "
     "an animal can trip filters tuned for human-facing harm. This observation directly shaped a design "
     "decision described later, where the content-filter thresholds were relaxed so that genuine medical "
     "information would not be silently discarded.")

h2("2.9 Web Technologies for AI Applications")
body("Delivering a model to real users requires more than the model itself. A typical modern stack separates "
     "a backend that exposes a REST API from a frontend single-page application that consumes it. Django and "
     "Django REST Framework provide a mature backend with an object-relational mapper, authentication, and "
     "serialisation, while React with a build tool such as Vite provides a responsive interface. This "
     "separation, used throughout the project, allows the model, the database, and the user interface to "
     "evolve independently, and it allows the computationally heavy work to stay on the server while the "
     "client remains lightweight.")

h2("2.10 Authentication and Security")
body("Any application that stores personal records needs a sound authentication scheme. Token-based "
     "authentication using JSON Web Tokens has become a common choice for single-page applications because "
     "it avoids server-side session storage and works naturally across separate frontend and backend "
     "deployments. A signed token carries the user's identity and an expiry time; the server verifies the "
     "signature on each request without needing to look the session up in a database. Established good "
     "practice, followed here, includes hashing passwords rather than storing them, keeping access tokens "
     "short-lived, restricting cross-origin requests, and never exposing third-party API keys to the client.")

h2("2.11 Summary and Research Gap")
body("The literature provides mature, well-understood building blocks: efficient convolutional architectures "
     "for constrained environments, transfer learning for small datasets, interpretable symptom matching, "
     "constrained generation from language models, and conventional secure web architecture. What is less "
     "commonly documented is the integration of all of these into a single, dependable application for a "
     "developing-country livestock context, together with an honest account of the engineering required to "
     "make the generative component reliable. This project addresses that gap.")
page_break()

# =====================================================================
# CHAPTER 3: SYSTEM ARCHITECTURE AND DESIGN
# =====================================================================
h1("Chapter 3: System Architecture and Design")
h2("3.1 Design Goals and Requirements")
body("The design was driven by a small number of clear goals. The system had to be usable by a non-technical "
     "owner, which meant plain language and a small number of obvious actions. It had to keep sensitive "
     "assets—the database and the Gemini API key—entirely on the server. It had to remain useful "
     "even when an external service failed, degrading to local guidance rather than breaking. And it had to "
     "be modular, so that the image model, the symptom engine, and the generative layer could each be "
     "changed without disturbing the others. The functional and non-functional requirements that follow from "
     "these goals were summarised earlier in Table 1.1.")

h2("3.2 Overall System Architecture")
body("VetAI Diagnostics follows a decoupled, three-tier client-server design. The presentation tier is a "
     "React single-page application that runs in the browser. The application tier is a Django REST "
     "Framework backend that exposes a set of JSON endpoints and coordinates all of the intelligent "
     "behaviour. The data tier is a MySQL database together with the external Gemini service and the local "
     "TensorFlow model. The browser communicates only with the backend, over HTTP, authenticated with JSON "
     "Web Tokens.")
body("Keeping these concerns separate has a clear benefit. The frontend never talks to the Gemini API or the "
     "database directly; it only knows about the REST endpoints. This means the API key stays on the server, "
     "the database is never exposed to the public internet, and the model can be retrained or replaced "
     "without touching the interface. The backend itself is divided into three Django applications "
     "(accounts, animal, and ml_model) so that authentication, the diagnostic logic, and the machine-learning "
     "assets are cleanly separated.")
figure_placeholder("Figure 3.1: Overall three-tier system architecture")

h2("3.3 Use Case Analysis")
body("The system has effectively one human actor, the registered user, who may be a farmer, a "
     "para-veterinary worker, or a student. An external actor, the Gemini service, supplies medical content. "
     "The principal use cases are listed in Table 3.1.")
make_table(
    ["Use Case", "Actor", "Description"],
    [
        ["Register / Log in", "User", "Create an account and obtain an authenticated session"],
        ["Diagnose from image", "User", "Upload a photograph and receive a classified result"],
        ["Diagnose from symptoms", "User", "Select symptoms and receive ranked candidates"],
        ["View report", "User", "Read the full guidance for a diagnosis"],
        ["Ask the assistant", "User", "Pose a veterinary question in chat"],
        ["View history", "User", "Browse and inspect past diagnoses"],
        ["View dashboard", "User", "See summary statistics and charts"],
        ["Provide guidance", "Gemini", "Return structured disease information on request"],
    ],
    widths=[1.8, 1.2, 3.6],
    title="Table 3.1: Primary use cases",
)
figure_placeholder("Figure 3.2: Use-case diagram for the registered user")

h2("3.4 Image Diagnosis Pipeline")
h3("3.4.1 Upload and Validation")
body("When a user submits an image, the request reaches the DetectAnimalAPIView endpoint. Before any "
     "processing, the view checks that a file was actually provided, that it is no larger than ten "
     "megabytes, and that its extension is one of the accepted image formats. These checks reject obviously "
     "invalid input early and protect the server from wasteful work. The file is then written to a temporary "
     "location using a cross-platform temporary-file mechanism so that it can be read by the image loader "
     "without assuming a particular operating system.")
h3("3.4.2 Inference")
body("Prediction is handled by a singleton wrapper around the trained Keras model. Implementing it as a "
     "singleton matters: loading a deep-learning model from disk is expensive, so the model is loaded once "
     "and reused across requests rather than re-read every time. The image is loaded at 224 by 224 pixels "
     "and passed to the network. Crucially, normalisation is performed by a Rescaling layer inside the model "
     "rather than in the preprocessing code, which keeps the training and inference pipelines identical and "
     "avoids a class of subtle bugs in which the two stages disagree about how pixels are scaled.")
h3("3.4.3 Confidence Handling")
body("The network outputs a probability for each of the three classes; the highest is taken as the "
     "prediction, and its value becomes the confidence score. The design deliberately surfaces low "
     "confidence rather than hiding it. When the top probability falls below a threshold, the response "
     "carries a warning advising the user to seek professional confirmation, and the situation is logged so "
     "that recurring uncertainty can be investigated and the model retrained if necessary.")
h3("3.4.4 Enrichment and Storage")
body("The predicted label is normalised and checked against the set of valid diseases, after which the "
     "backend requests fresh medical information from Gemini. The combined result—prediction, "
     "confidence, and guidance—is saved as an AnimalDetection record linked to the user, the temporary "
     "file is removed, and a serialised report is returned to the frontend together with a marker indicating "
     "whether the guidance came from the live service or from the local fallback.")
figure_placeholder("Figure 3.3: Image diagnosis pipeline")

h2("3.5 Symptom Diagnosis Engine")
body("The symptom engine is implemented in a DatasetProcessor class that loads the training and testing CSV "
     "files once and caches the result, again using a singleton so that the comparatively large dataset is "
     "read from disk only once per process. Each row of the dataset is a one-hot vector over roughly "
     "ninety-three symptom columns together with a disease label in the prognosis column. From this, the "
     "processor builds a map from each disease to the symptoms associated with it, and an index of how "
     "frequently each symptom appears for each disease.")
body("When a user submits a set of symptoms, the engine normalises the symptom names and compares them "
     "against every disease, counting the matches using a tolerant substring comparison so that minor "
     "differences in wording do not prevent a match. It then computes a confidence value that blends two "
     "ideas: how many of the user's symptoms were matched, and how well those matches cover the disease's "
     "known symptoms. The formula used is confidence equal to 0.6 times the match ratio plus 0.4 times the "
     "coverage. The weighting favours matching the user's reported symptoms while still rewarding a close "
     "overall fit. The top five diseases are returned, ranked by this score, and the best match is passed to "
     "Gemini for enrichment with detailed guidance.")
figure_placeholder("Figure 3.4: Symptom diagnosis scoring and enrichment flow")

h2("3.6 Generative AI Integration Design")
body("The Gemini integration is the most carefully engineered part of the system, because a generative "
     "model that fails silently would undermine the whole experience. The service builds a strict prompt "
     "that asks for a fixed JSON structure containing the disease name, severity, symptoms, treatment, "
     "prevention, contagion flag, and antibiotics. It requests a JSON response type and supplies safety "
     "settings, and it lowers the sampling temperature so that the output is stable.")
body("Reliability is achieved through several layers, which were strengthened substantially during the "
     "project. The service tries more than one Gemini model and both the v1beta and v1 API versions. A "
     "response that returns an HTTP 200 status but contains no usable text—for example, a "
     "safety-blocked or empty candidate—is no longer accepted; instead the service detects the empty "
     "body and moves on to the next model, so that one blocked response cannot abort the whole request. "
     "Truncated or malformed JSON is repaired where possible by a dedicated helper, and if the model still "
     "returns an incomplete object the request is retried with a more compact prompt that is less likely to "
     "be cut off. Only when every attempt fails does the system fall back to a local JSON knowledge base, "
     "and finally to a safe default that simply advises consulting a veterinarian. As part of this work the "
     "content-filter thresholds were relaxed from the default of blocking medium-and-above to blocking only "
     "high-severity content, which noticeably reduced the number of legitimate veterinary responses that "
     "were being discarded.")
figure_placeholder("Figure 3.5: Gemini request, retry, and fallback logic")

h2("3.7 Database Design")
body("Two main tables hold the system's data, both linked to the built-in user table by a foreign key. The "
     "AnimalDetection table stores image-based results: the uploaded image, the predicted disease, the "
     "confidence, the severity, and the guidance returned for that case, along with the full set of "
     "per-class probabilities. The SymptomDiagnosis table stores symptom-based results, including the "
     "symptoms the user entered, the symptoms that were matched, the match rate, and the complete list of "
     "ranked candidates retained for reference. List-valued fields such as symptoms, treatment steps, and "
     "antibiotics are stored as JSON columns, which keeps a complete diagnosis in a single row and avoids a "
     "proliferation of small related tables. Both tables record a status and creation and update "
     "timestamps, and both are indexed on the user and the creation time so that a user's history can be "
     "retrieved efficiently in reverse chronological order.")
figure_placeholder("Figure 3.6: Entity-relationship diagram of the database")

h2("3.8 Authentication Design")
body("Access to every diagnostic feature requires a logged-in user. Authentication uses JSON Web Tokens "
     "issued by the SimpleJWT library. On successful login the server returns a short-lived access token and "
     "a longer-lived refresh token; the frontend stores these in the browser and attaches the access token "
     "to every subsequent request as a bearer token. On the client side, a route guard prevents the main "
     "application from rendering at all unless a token is present, and an expired or invalid token causes the "
     "user to be redirected back to the sign-in page. On the server side, every protected view declares that "
     "it requires authentication and filters its data by the requesting user, so the records of one user are "
     "never visible to another.")

h2("3.9 End-to-End Data Flow")
body("Taken together, a typical interaction flows as follows. The user logs in and receives tokens. They "
     "choose a feature and submit input. The frontend sends that input, with the access token, to the "
     "appropriate REST endpoint. The backend authenticates the request, runs the relevant intelligence—"
     "the model, the symptom engine, or the chat proxy—enriches the result with Gemini where "
     "appropriate, stores any record that needs to persist, and returns a JSON response. The frontend then "
     "renders that response as a report, a chart, or a chat reply. No step in this chain requires the client "
     "to hold a secret or to know anything about the database.")
page_break()

# =====================================================================
# CHAPTER 4: IMPLEMENTATION
# =====================================================================
h1("Chapter 4: Implementation")
h2("4.1 Technology Stack")
body("The complete set of technologies used in the project is summarised in Table 4.1. Versions reflect "
     "those pinned in the project's requirements and package files.")
make_table(
    ["Layer", "Technology", "Version", "Purpose"],
    [
        ["Frontend", "React", "19.1", "Component-based user interface"],
        ["", "Vite", "7.1", "Dev server and production build"],
        ["", "Tailwind CSS", "4.1", "Utility-first styling"],
        ["", "TanStack React Query", "5.90", "Server state and mutations"],
        ["", "Axios", "1.13", "HTTP client"],
        ["", "Recharts", "3.5", "Dashboard charts"],
        ["", "lucide-react", "0.5x", "Icon set"],
        ["Backend", "Django", "5.2", "Web framework and ORM"],
        ["", "Django REST Framework", "3.16", "REST API layer"],
        ["", "SimpleJWT", "5.3", "JWT authentication"],
        ["", "django-cors-headers", "4.7", "Cross-origin handling"],
        ["AI / ML", "TensorFlow / Keras", "2.13", "CNN training and inference"],
        ["", "MobileNetV2", "ImageNet", "Transfer-learning backbone"],
        ["", "scikit-learn", "1.7", "Metrics and class weights"],
        ["", "pandas / NumPy", "2.2 / 1.x", "Dataset and array processing"],
        ["", "Pillow / OpenCV", "10.3 / 4.8", "Image loading and processing"],
        ["", "Google Gemini API", "2.5-flash", "Real-time guidance and chat"],
        ["Database", "MySQL", "8.x", "Users and diagnosis records"],
        ["Server", "Gunicorn / WhiteNoise", "23 / 6.8", "WSGI server / static files"],
        ["Hosting", "Railway / Vercel", "-", "Backend / frontend hosting"],
    ],
    widths=[1.0, 2.3, 1.1, 2.2],
    title="Table 4.1: Technology stack used in the project",
)

h2("4.2 Development Environment and Tools")
body("Development was carried out on Windows using Visual Studio Code as the editor, with Git for version "
     "control. Python virtual environments isolated the backend dependencies, and Node.js with npm managed "
     "the frontend packages. The Django development server and the Vite development server were run side by "
     "side during work, with the frontend configured to call the local backend. One small but useful detail "
     "was a wrapper around Python's print function in the backend that strips non-ASCII characters, because "
     "the emoji used in log messages would otherwise crash the default Windows console encoding.")

h2("4.3 Backend Implementation")
h3("4.3.1 Project Structure")
body("The Django project is divided into three applications. The accounts application handles user "
     "registration. The animal application contains the core logic: the API views, the database models, the "
     "Gemini service, a JSON-repair helper, and the dataset processor. The ml_model application holds the "
     "training code, the local disease knowledge base, the management command used to train the model, and "
     "the saved model files. Project-wide settings, the root URL configuration, and the WSGI and ASGI entry "
     "points live in a separate settings package.")
figure_placeholder("Figure 4.1: Backend application structure")
h3("4.3.2 Configuration Management")
body("Configuration is read from environment variables through a dotenv loader, so that secrets such as the "
     "database password and the Gemini API key are never committed to source control. Sensible defaults are "
     "provided for local development, while production values are supplied by the hosting platform. The same "
     "mechanism controls debug mode, the allowed hosts, and the cross-origin whitelist.")
h3("4.3.3 Data Models")
body("The models are defined in the animal application. As described in the design chapter, AnimalDetection "
     "and SymptomDiagnosis both use JSON fields to store list-valued data such as symptoms, treatment steps, "
     "and antibiotics. Each model also defines its allowed status values and the indexes used to speed up "
     "history queries.")
h3("4.3.4 Detection View")
body("DetectAnimalAPIView ties the image pipeline together. It performs validation, runs the model, requests "
     "guidance from Gemini with caching disabled so that the information is always current, handles the "
     "fallback chain, stores the record, and returns the serialised result. The view distinguishes between "
     "validation errors, missing-model errors, and unexpected failures, returning an appropriate HTTP status "
     "and a human-readable message in each case rather than leaking a stack trace to the client.")
h3("4.3.5 Gemini Service")
body("The gemini_service module is responsible for all communication with the Gemini API. It builds the "
     "prompt, sends the request across candidate models and API versions, parses and repairs the JSON, and "
     "applies sensible defaults so that the user interface is never handed an empty field. Error messages "
     "are sanitised so that the API key can never leak into logs. The reliability improvements described in "
     "the design chapter—validating that a 200 response actually contains text, skipping empty or "
     "blocked responses, retrying with backoff, and relaxing the safety thresholds—are implemented "
     "here.")
h3("4.3.6 Dataset Processor")
body("The dataset_processor module loads the CSV dataset, normalises symptom names, builds the disease-to-"
     "symptom map, and exposes the scoring method used for symptom diagnosis. It also provides the list of "
     "available symptoms that the frontend displays for selection, and a method that returns detailed "
     "information for a single disease.")
h3("4.3.7 Vet Chat Proxy")
body("The chat assistant is implemented as a server-side proxy. The view builds a veterinary system prompt "
     "that instructs the model to answer professionally, to avoid definitive diagnoses, and to recommend "
     "professional care for serious symptoms. It appends a bounded slice of the recent conversation history "
     "and forwards the request to Gemini, trying fallback models if the primary one is unavailable. Because "
     "the call is made from the server, the API key is never exposed to the browser.")

h2("4.4 Model Training")
body("The training code is contained in the ml_model application and is invoked through a Django management "
     "command, which means the model can be retrained with a single command from within the project. The "
     "network is assembled as described below and summarised in Table 4.2.")
make_table(
    ["Layer", "Configuration", "Output"],
    [
        ["Input", "224 x 224 x 3 RGB image", "224 x 224 x 3"],
        ["Rescaling", "Divide pixel values by 255", "224 x 224 x 3"],
        ["MobileNetV2 base", "ImageNet weights, frozen", "7 x 7 x 1280 features"],
        ["GlobalAveragePooling2D", "Spatial average", "1280"],
        ["Dense + ReLU", "256 units", "256"],
        ["Dropout", "Rate 0.5", "256"],
        ["Dense + ReLU", "128 units", "128"],
        ["Dropout", "Rate 0.3", "128"],
        ["Dense + Softmax", "n class units", "n probabilities"],
    ],
    widths=[2.2, 2.6, 1.8],
    title="Table 4.2: CNN layer configuration",
)
figure_placeholder("Figure 4.2: Convolutional neural network architecture")
body("Several decisions during training are worth noting. Because the dataset is not perfectly balanced "
     "across the classes, class weights are computed automatically and supplied to the training loop so that "
     "rarer classes are not overwhelmed by common ones. Data augmentation—rotation, shifts, horizontal "
     "flips, zoom, and shear—is applied to the training images to improve generalisation, but, "
     "importantly, the augmentation does not rescale the pixels, because rescaling is handled inside the "
     "model; this keeps the training and inference preprocessing consistent. Three callbacks guide the "
     "training: early stopping restores the best weights if validation loss stops improving, a learning-rate "
     "reducer lowers the rate when progress plateaus, and a checkpoint saves the best model by validation "
     "accuracy. After training, the model is saved in both the HDF5 and TensorFlow SavedModel formats, and "
     "the class index map is written as both a pickle and a JSON file, so that inference can load whichever "
     "format the runtime supports.")

h2("4.5 Frontend Implementation")
h3("4.5.1 Application Structure")
body("The React application is organised around a protected home page that hosts a tabbed dashboard. The "
     "tabs correspond to the main features, and state that needs to be shared—such as the most recent "
     "diagnosis—is held in the home component so that a successful diagnosis can automatically switch "
     "the user to the report tab. Routing is handled by react-router-dom, with a dedicated route guard "
     "component protecting the application shell.")
figure_placeholder("Figure 4.3: Frontend component hierarchy")
make_table(
    ["Component", "Responsibility"],
    [
        ["SignIn / SignUp", "Authentication forms with validation"],
        ["ProtectedRoute", "Blocks the app unless a token is present"],
        ["Home", "Tabbed shell and shared state"],
        ["Dashboard", "Statistics and charts via Recharts"],
        ["ImageDiagnosis", "Image upload and prediction request"],
        ["SymptomDiagnosis", "Symptom selection and diagnosis"],
        ["DiseaseResults", "Full clinical report rendering"],
        ["DetectionHistory", "List and detail of past records"],
        ["VetChat", "Conversational assistant UI"],
        ["LogOut", "Clears tokens and redirects"],
    ],
    widths=[2.0, 4.6],
    title="Table 4.4: Frontend components and responsibilities",
)
h3("4.5.2 Data Fetching")
body("Network calls use Axios, wrapped by TanStack React Query. Queries load data such as the list of "
     "available symptoms, while mutations handle actions such as uploading an image or submitting symptoms. "
     "React Query manages the loading, error, and success states, which keeps the components themselves "
     "focused on presentation. A shared configuration module supplies the API base URL and attaches the "
     "bearer token, so that authentication is handled consistently across every request.")
h3("4.5.3 Report and Charts")
body("The diagnosis report component renders the severity, symptoms, treatment protocol, prevention "
     "measures, and recommended antibiotics, and shows a prominent warning when a condition is contagious. "
     "If any guidance field is missing, the component requests the missing details from a disease-detail "
     "endpoint, which makes the report robust to partial data. The dashboard uses Recharts to display "
     "totals, average confidence, the distribution of diseases as a pie chart, a severity breakdown, and a "
     "trend of recent activity.")
h3("4.5.4 Chat Interface")
body("The chat component stores its conversation in the browser's local storage so that history survives a "
     "page reload, presents quick-question buttons for common queries, and shows an animated typing "
     "indicator while the assistant is responding. Errors from the server are translated into friendly "
     "messages rather than raw status codes.")

h2("4.6 Authentication System")
h3("4.6.1 JWT Token Flow")
body("Registration posts a username, email, and password to the accounts endpoint, where Django creates the "
     "user and stores a hashed password. Login posts credentials to the token endpoint, which returns an "
     "access token valid for twenty-four hours and a refresh token valid for seven days, with rotation "
     "enabled. The frontend stores both and sends the access token as a bearer token on every request.")
figure_placeholder("Figure 4.4: JWT authentication token flow")
h3("4.6.2 Security Details")
bullet("Passwords are hashed by Django's authentication system using a strong algorithm; plain-text "
       "passwords are never stored.")
bullet("Every protected endpoint requires authentication and filters data by the requesting user, so one "
       "user can never see another user's records.")
bullet("The Gemini API key is held only on the server, and error text is sanitised to remove it before "
       "logging.")
bullet("Cross-origin requests are restricted to the expected frontend origins, and uploads are validated for "
       "size and type before processing.")

h2("4.7 API Design")
body("The REST API is summarised in Table 4.3. All endpoints except registration and token issuance require "
     "a valid bearer token. Responses follow a consistent shape, typically a success flag together with a "
     "data payload, which simplifies handling on the client.")
make_table(
    ["Endpoint", "Method", "Auth", "Description"],
    [
        ["/api/register/", "POST", "No", "Create a user account"],
        ["/api/token/", "POST", "No", "Obtain access and refresh tokens"],
        ["/api/token/refresh/", "POST", "No", "Refresh the access token"],
        ["/api/animal/detect/", "POST", "Yes", "Image-based disease detection"],
        ["/api/animal/history/", "GET", "Yes", "List image detections"],
        ["/api/animal/detail/<id>/", "GET / PATCH", "Yes", "Retrieve or update a detection"],
        ["/api/animal/statistics/", "GET", "Yes", "Per-user analytics"],
        ["/api/symptoms/", "GET", "Yes", "List available symptoms"],
        ["/api/diagnose/", "POST", "Yes", "Symptom-based diagnosis"],
        ["/api/diseases/<id>/", "GET", "Yes", "Detailed disease information"],
        ["/api/animal/vetchat/", "POST", "Yes", "Veterinary chat assistant"],
        ["/health/", "GET", "No", "Service health check"],
    ],
    widths=[2.3, 1.1, 0.7, 2.5],
    title="Table 4.3: REST API endpoints",
)

h2("4.8 Deployment")
body("The backend is deployed to Railway and served by Gunicorn, with WhiteNoise handling static files so "
     "that no separate web server is needed for them. A Procfile declares the start command and a runtime "
     "file pins the Python version. The frontend is built by Vite into static assets and hosted on Vercel, "
     "with the API base URL supplied as a build-time environment variable. The cross-origin and trusted-"
     "origin settings on the backend whitelist the deployed frontend domain, and secure-cookie settings are "
     "enabled automatically when debug mode is off.")
page_break()

# =====================================================================
# CHAPTER 5: RESULTS AND DISCUSSION
# =====================================================================
h1("Chapter 5: Results and Discussion")
h2("5.1 Trained Model and Supported Diseases")
body("The image classifier was trained on a dataset organised into folders by class. After training, the "
     "saved class mapping confirms that the model recognises three categories, listed in Table 5.1. The "
     "model is stored in three forms—an HDF5 file, a TensorFlow SavedModel directory, and a pickled "
     "class index—so that it can be loaded under different runtime configurations.")
make_table(
    ["Class Index", "Label", "Disease"],
    [
        ["0", "foot-and-mouth", "Foot and Mouth Disease"],
        ["1", "healthy", "Healthy animal (no disease)"],
        ["2", "lumpy", "Lumpy Skin Disease"],
    ],
    widths=[1.4, 2.4, 2.8],
    title="Table 5.1: Trained model classes and supported diseases",
)

h2("5.2 Testing Methodology")
body("Testing was carried out manually and feature by feature, in line with the iterative development "
     "approach. Each endpoint was exercised both through the user interface and directly, to confirm that it "
     "behaved correctly for valid input and failed gracefully for invalid input. Particular attention was "
     "paid to the boundary conditions identified during design: oversized files, unsupported formats, empty "
     "symptom lists, expired tokens, and an unavailable Gemini service. The aim was not a formal test suite "
     "with coverage metrics, which was beyond the scope of a single-developer project, but a thorough "
     "demonstration that each requirement from Table 1.1 was met.")

h2("5.3 Functional Testing Results")
body("Table 5.2 lists representative functional test cases and their outcomes.")
make_table(
    ["ID", "Test Case", "Expected Result", "Outcome"],
    [
        ["T-1", "Register with valid details", "Account created", "Pass"],
        ["T-2", "Log in with correct credentials", "Tokens issued, redirect to dashboard", "Pass"],
        ["T-3", "Access app without a token", "Redirect to sign-in", "Pass"],
        ["T-4", "Upload a valid image", "Prediction with confidence and report", "Pass"],
        ["T-5", "Upload a file over 10 MB", "Rejected with a clear message", "Pass"],
        ["T-6", "Upload an unsupported format", "Rejected with a clear message", "Pass"],
        ["T-7", "Submit selected symptoms", "Ranked candidate diseases returned", "Pass"],
        ["T-8", "Submit no symptoms", "Validation message, no diagnosis", "Pass"],
        ["T-9", "Open the dashboard", "Statistics and charts rendered", "Pass"],
        ["T-10", "Ask the chat assistant a question", "Professional, relevant reply", "Pass"],
        ["T-11", "Use the app with Gemini unavailable", "Local fallback guidance shown", "Pass"],
    ],
    widths=[0.6, 2.6, 2.4, 0.9],
    title="Table 5.2: Functional test cases and outcomes",
)

h2("5.4 System Behaviour")
body("End to end, a typical image diagnosis completes within a few seconds: the largest single cost is the "
     "real-time Gemini call, since the model inference itself is fast once the model is loaded. The "
     "application correctly stores each result, updates the dashboard statistics, and shows the report. The "
     "symptom path behaves similarly, returning a ranked list of candidate diseases together with the "
     "symptoms that were matched for each.")
body("During testing, the most common source of trouble was not the model's accuracy but the loading of the "
     "model across TensorFlow versions, because the file had been saved under one version and was being "
     "loaded under another. This produced errors related to layer configuration that did not appear during "
     "training. The problem was handled by a defensive loader that normalises legacy configuration keys and "
     "supplies compatible replacements for the affected layers, after which the model loads consistently. "
     "The episode is a reminder that reproducibility across environments deserves as much attention as model "
     "accuracy itself.")

h2("5.5 Reliability of the Generative Layer")
body("A specific goal during development was to reduce the cases where the symptom diagnosis returned no "
     "real-time guidance. Investigation showed three causes: the content-safety filters were occasionally "
     "blocking legitimate veterinary text; a response that returned an HTTP 200 status with an empty or "
     "blocked body was being treated as success, after which parsing failed; and only a small number of "
     "retries were attempted. Table 5.3 summarises the changes made and their effect.")
make_table(
    ["Change", "Before", "After"],
    [
        ["Safety filter threshold", "Block medium and above", "Block only high"],
        ["Empty / blocked 200 response", "Accepted, then parse failed", "Skipped; next model tried"],
        ["Retry attempts (with backoff)", "2", "3"],
        ["Net effect", "Guidance often missing", "Guidance returned in most cases"],
    ],
    widths=[2.6, 2.2, 2.2],
    title="Table 5.3: Gemini reliability improvements",
)
body("Taken together, these changes mean that a single blocked or empty response can no longer cause the "
     "whole request to fall back to generic text. The service now exhausts several model and version "
     "combinations, each a genuine attempt at obtaining a usable answer, before resorting to the local "
     "knowledge base. After these changes, real-time guidance was returned in the large majority of test "
     "queries, where previously it had frequently been replaced by the generic fallback.")
figure_placeholder("Figure 5.1: Confidence distribution across sample predictions")

h2("5.6 Performance Considerations")
body("The system was designed to be light on resources. The model is loaded once and cached, so only the "
     "first request after a restart pays the loading cost. Image preprocessing and inference are quick on a "
     "single image. The dominant latency is the external Gemini call, which is bounded by a timeout and "
     "mitigated by trying faster models first. On the database side, the indexes on the user and creation "
     "time keep history and statistics queries fast even as records accumulate. Where stricter latency is "
     "required in future, the existing but currently disabled response cache for disease information could be "
     "enabled for non-critical lookups.")

h2("5.7 Security Assessment")
body("Measured against the security requirements, the system performs well in the areas that matter most "
     "for a student project: passwords are hashed, tokens are signed and time-limited, the API key never "
     "leaves the server, cross-origin access is restricted, and uploads are validated. Per-user data "
     "isolation was verified by confirming that one account cannot retrieve another account's records. Some "
     "hardening remains for a production deployment, notably ensuring that the secret key and database "
     "credentials are always supplied through the environment rather than relying on development defaults, "
     "and adding rate limiting to the authentication and detection endpoints. These are documented honestly "
     "as limitations rather than presented as solved.")

h2("5.8 User Interface Assessment")
body("Informal use of the interface by classmates suggested that the tabbed layout was easy to follow and "
     "that the diagnosis report communicated the key facts clearly. The contagion warning, rendered as a "
     "prominent coloured panel, drew attention as intended. The scanning animation shown while an image is "
     "being analysed helped reassure users that the system was working rather than stalled. The most common "
     "request was for the ability to export a report, which has been recorded as future work.")
figure_placeholder("Figure 5.2: Application screenshots (dashboard, diagnosis, report)")

h2("5.9 Comparison with Existing Approaches")
body("Table 5.4 positions the system against the kinds of tools commonly available. The comparison is "
     "qualitative, since the alternatives differ widely in purpose, but it illustrates where the combined "
     "approach adds value.")
make_table(
    ["Capability", "Generic web search", "Single-purpose classifier", "VetAI Diagnostics"],
    [
        ["Image-based detection", "No", "Yes", "Yes"],
        ["Symptom-based diagnosis", "Partial", "No", "Yes"],
        ["Real-time tailored guidance", "No", "No", "Yes"],
        ["Contagion warning", "No", "Rarely", "Yes"],
        ["Per-user history", "No", "Rarely", "Yes"],
        ["Graceful offline fallback", "n/a", "Varies", "Yes"],
    ],
    widths=[2.2, 1.6, 1.7, 1.5],
    title="Table 5.4: Comparison with existing approaches",
)

h2("5.10 Challenges Faced")
body("Three challenges stood out during the project. The first was the cross-version model-loading problem "
     "described above, which consumed a disproportionate amount of debugging time and ultimately required a "
     "carefully layered loader. The second was the unreliability of the generative layer, which was solved "
     "by the validation, retry, and safety-threshold changes documented earlier. The third was simpler but "
     "persistent: keeping the two diagnostic paths, with their different disease vocabularies, coherent in "
     "the user interface, which was addressed by making the report component tolerant of partial data and "
     "able to fetch missing details on demand.")
page_break()

# =====================================================================
# CHAPTER 6: CONCLUSIONS AND FUTURE WORK
# =====================================================================
h1("Chapter 6: Conclusions and Future Work")
h2("6.1 Summary of Contributions")
body("This project contributes a working, end-to-end diagnostic aid for cattle disease that combines three "
     "distinct forms of artificial intelligence within a conventional secure web application. Its specific "
     "contributions are a transfer-learning image classifier for three cattle conditions, an interpretable "
     "symptom-scoring engine over a tabular dataset, and—most significantly from an engineering "
     "standpoint—a hardened generative-AI integration that returns reliable, structured medical "
     "guidance and degrades gracefully when the external service fails.")

h2("6.2 Conclusions")
body("The thesis set out to build an accessible diagnostic aid for cattle disease, and the resulting system "
     "meets that aim. VetAI Diagnostics accepts either an image or a set of symptoms, returns a probable "
     "diagnosis with a confidence score, and enriches every result with real-time guidance on treatment, "
     "prevention, and medication. It does so behind a secure, authenticated REST API, stores a per-user "
     "history, and presents the information through a clean single-page interface.")
body("Beyond assembling the components, a meaningful part of the work was making the generative layer "
     "dependable. The improvements to the Gemini integration—validating responses, trying multiple "
     "models, relaxing overly strict filters, and retrying sensibly—turned an occasionally unreliable "
     "feature into one that returns useful guidance in the large majority of cases. The broader lesson is "
     "that integrating a generative model into a real product is as much about defensive engineering around "
     "the model as it is about the model itself.")

h2("6.3 Limitations")
bullet("The image model recognises only three conditions; anything outside this set cannot be detected from "
       "a photograph.")
bullet("The two diagnostic paths use different disease vocabularies, because the image model and the symptom "
       "dataset were built from different sources.")
bullet("The local fallback knowledge base covers only a handful of diseases, so if the Gemini API is "
       "unavailable for a disease outside that set, only generic guidance is shown.")
bullet("The defensive model loader, while effective, is complex, and a cleaner solution would be to pin a "
       "single TensorFlow version and retrain.")
bullet("Evaluation was informal and based on a limited number of test cases and users rather than a large, "
       "labelled benchmark or a formal user study.")
bullet("Some production hardening, such as rate limiting and mandatory environment-supplied secrets, remains "
       "to be done.")

h2("6.4 Future Work")
bullet("Extend the image model to more diseases and to additional species such as goats and sheep, with a "
       "correspondingly larger and more balanced image dataset.")
bullet("Add Urdu-language support, both in the interface and in the generated guidance, to widen access "
       "among the intended users.")
bullet("Expand the local knowledge base so that the fallback is genuinely useful for every disease in the "
       "symptom dataset, not only the handful currently covered.")
bullet("Introduce response caching for non-critical lookups to reduce latency and API usage.")
bullet("Develop a native mobile application, or a progressive web app, so that owners can diagnose directly "
       "in the field with limited connectivity.")
bullet("Add the ability to export a diagnosis as a PDF report for sharing with a veterinarian.")
bullet("Conduct a formal evaluation, including a labelled image test set for the classifier and a structured "
       "user study for the interface.")

h2("6.5 Final Remarks")
body("VetAI Diagnostics began as an attempt to put recent advances in computer vision and generative AI to "
     "practical use for a real and local problem. The finished system is modest in scope but complete in "
     "execution, and the engineering lessons learned—particularly around making external AI services "
     "dependable—are likely to outlast the specific application. With the extensions outlined above, "
     "the same foundation could grow into a genuinely useful tool for livestock owners.")
page_break()

# =====================================================================
# REFERENCES
# =====================================================================
h1("References")
refs = [
    "M. Sandler, A. Howard, M. Zhu, A. Zhmoginov, and L.-C. Chen, “MobileNetV2: Inverted Residuals and "
    "Linear Bottlenecks,” in Proc. IEEE Conf. Computer Vision and Pattern Recognition (CVPR), 2018, "
    "pp. 4510–4520.",
    "A. G. Howard et al., “MobileNets: Efficient Convolutional Neural Networks for Mobile Vision "
    "Applications,” arXiv preprint arXiv:1704.04861, 2017.",
    "K. He, X. Zhang, S. Ren, and J. Sun, “Deep Residual Learning for Image Recognition,” in Proc. "
    "IEEE CVPR, 2016, pp. 770–778.",
    "K. Simonyan and A. Zisserman, “Very Deep Convolutional Networks for Large-Scale Image "
    "Recognition,” in Proc. Int. Conf. Learning Representations (ICLR), 2015.",
    "A. Krizhevsky, I. Sutskever, and G. E. Hinton, “ImageNet Classification with Deep Convolutional "
    "Neural Networks,” in Advances in Neural Information Processing Systems (NeurIPS), 2012.",
    "J. Deng, W. Dong, R. Socher, L.-J. Li, K. Li, and L. Fei-Fei, “ImageNet: A Large-Scale "
    "Hierarchical Image Database,” in Proc. IEEE CVPR, 2009, pp. 248–255.",
    "S. J. Pan and Q. Yang, “A Survey on Transfer Learning,” IEEE Trans. Knowledge and Data "
    "Engineering, vol. 22, no. 10, pp. 1345–1359, 2010.",
    "S. P. Mohanty, D. P. Hughes, and M. Salathe, “Using Deep Learning for Image-Based Plant Disease "
    "Detection,” Frontiers in Plant Science, vol. 7, art. 1419, 2016.",
    "M. Abadi et al., “TensorFlow: A System for Large-Scale Machine Learning,” in Proc. 12th USENIX "
    "Symp. Operating Systems Design and Implementation (OSDI), 2016, pp. 265–283.",
    "F. Chollet, Deep Learning with Python, 2nd ed. Shelter Island, NY: Manning Publications, 2021.",
    "P. Lewis et al., “Retrieval-Augmented Generation for Knowledge-Intensive NLP Tasks,” in "
    "Advances in Neural Information Processing Systems (NeurIPS), 2020.",
    "T. Brown et al., “Language Models are Few-Shot Learners,” in Advances in Neural Information "
    "Processing Systems (NeurIPS), 2020.",
    "Google, “Gemini API Documentation,” Google AI for Developers. [Online]. Available: "
    "https://ai.google.dev/",
    "Django Software Foundation, “Django Documentation,” Version 5.2. [Online]. Available: "
    "https://docs.djangoproject.com/",
    "Encode OSS Ltd., “Django REST Framework Documentation.” [Online]. Available: "
    "https://www.django-rest-framework.org/",
    "Meta Open Source, “React Documentation.” [Online]. Available: https://react.dev/",
    "M. Jones, J. Bradley, and N. Sakimura, “JSON Web Token (JWT),” RFC 7519, Internet Engineering "
    "Task Force, 2015.",
    "World Organisation for Animal Health (WOAH), “Foot and Mouth Disease,” Technical Disease "
    "Card. [Online]. Available: https://www.woah.org/",
    "World Organisation for Animal Health (WOAH), “Lumpy Skin Disease,” Technical Disease Card. "
    "[Online]. Available: https://www.woah.org/",
    "Food and Agriculture Organization of the United Nations (FAO), “The Future of Livestock in "
    "Pakistan,” FAO Report, 2018.",
]
for i, r in enumerate(refs, 1):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.4)
    p.paragraph_format.first_line_indent = Inches(-0.4)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.add_run(f"[{i}]\t{r}")
page_break()

# =====================================================================
# APPENDICES
# =====================================================================
h1("Appendices")
h2("Appendix A: Installation and Setup Guide")
h3("Prerequisites")
bullet("Python 3.10 or higher")
bullet("Node.js 18 or higher")
bullet("MySQL 8.x")
bullet("A Google Gemini API key")
h3("Backend Setup")
body("From the server directory, create and activate a virtual environment, install the dependencies with "
     "pip install -r requirements.txt, set the environment variables described in Appendix B, apply the "
     "database migrations with python manage.py migrate, optionally train the model with python manage.py "
     "train_model, and start the development server with python manage.py runserver 8000.")
code_block([
    "cd server",
    "python -m venv env",
    "env\\Scripts\\activate        # Windows",
    "pip install -r requirements.txt",
    "python manage.py migrate",
    "python manage.py runserver 8000",
])
h3("Frontend Setup")
body("From the client directory, install dependencies with npm install, set VITE_API_BASE_URL to the address "
     "of the backend API, and start the development server with npm run dev. The application will be "
     "available at the local Vite address, by default on port 5173.")
code_block([
    "cd client",
    "npm install",
    "# create .env with VITE_API_BASE_URL=http://127.0.0.1:8000/api",
    "npm run dev",
])

h2("Appendix B: Environment Variables Reference")
make_table(
    ["Variable", "Purpose"],
    [
        ["SECRET_KEY", "Django secret key (must be set in production)"],
        ["DEBUG", "Set to False in production"],
        ["DB_NAME / DB_USER / DB_PASSWORD / DB_HOST / DB_PORT", "MySQL connection settings"],
        ["GEMINI_API_KEY", "Key for the Google Gemini API"],
        ["GEMINI_MODEL", "Primary Gemini model (default gemini-2.5-flash)"],
        ["GEMINI_FALLBACK_MODELS", "Comma-separated list of fallback models"],
        ["CORS_ALLOWED_ORIGINS", "Additional allowed frontend origins"],
        ["VITE_API_BASE_URL", "Frontend setting pointing to the backend API"],
    ],
    widths=[3.2, 3.4],
    title="Table B.1: Environment variables reference",
)

h2("Appendix C: Database Schema")
body("Table C.1 lists the principal fields of the AnimalDetection model, and Table C.2 lists those of the "
     "SymptomDiagnosis model.")
make_table(
    ["Field", "Type", "Description"],
    [
        ["user", "Foreign key", "Owner of the record"],
        ["image", "Image", "Uploaded animal photograph"],
        ["animal_name", "Text", "Animal label (e.g. Cow)"],
        ["disease_name", "Text", "Predicted disease"],
        ["confidence_score", "Float", "Prediction confidence (0–1)"],
        ["severity", "Choice", "None / Low / Medium / High / Critical"],
        ["symptoms / treatment / prevention / antibiotics", "JSON list", "Guidance fields"],
        ["contagious", "Boolean", "Whether the disease is contagious"],
        ["all_predictions", "JSON", "Per-class probabilities"],
        ["status", "Choice", "diagnosed / treated / recovered / pending"],
        ["created_at / updated_at", "Datetime", "Timestamps"],
    ],
    widths=[2.7, 1.3, 2.6],
    title="Table C.1: AnimalDetection database fields",
)
make_table(
    ["Field", "Type", "Description"],
    [
        ["user", "Foreign key", "Owner of the record"],
        ["animal_name / animal_age", "Text / Integer", "Optional animal details"],
        ["input_symptoms", "JSON list", "Symptoms entered by the user"],
        ["disease_name / disease_id", "Text", "Best-match disease"],
        ["confidence_score / match_rate", "Float", "Score and proportion of symptoms matched"],
        ["matched_symptoms", "JSON list", "Symptoms that matched the disease"],
        ["treatment / prevention / medicines", "JSON list", "Guidance fields"],
        ["contagious", "Boolean", "Whether the disease is contagious"],
        ["all_results", "JSON", "Full ranked list of candidates"],
        ["status", "Choice", "diagnosed / treated / recovered / pending"],
        ["created_at / updated_at", "Datetime", "Timestamps"],
    ],
    widths=[2.7, 1.3, 2.6],
    title="Table C.2: SymptomDiagnosis database fields",
)

h2("Appendix D: Sample Disease Knowledge-Base Entry")
body("The local fallback knowledge base stores disease information as JSON. A representative entry for Lumpy "
     "Skin Disease is shown below.")
code_block([
    '"lumpy": {',
    '  "name": "Lumpy Skin Disease",',
    '  "severity": "High",',
    '  "symptoms": ["Fever", "Skin nodules on body", "Swollen lymph nodes", ...],',
    '  "treatment": ["Supportive care", "Antibiotics for secondary infection", ...],',
    '  "prevention": ["Vaccination", "Vector control", "Biosecurity", ...],',
    '  "contagious": true,',
    '  "antibiotics": ["Oxytetracycline", "Penicillin", "Streptomycin"]',
    '}',
])

doc.save(r"e:\HOme\FinalFYP\VetAI_Diagnostics_FYP_Report.docx")
print("SAVED: VetAI_Diagnostics_FYP_Report.docx")
print("paragraphs:", len(doc.paragraphs), "tables:", len(doc.tables))
